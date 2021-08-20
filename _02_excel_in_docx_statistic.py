# pip install python-docx
# pip install numpy
# pip install pandas
# 只能针对单一类型的文件，几乎没有泛化性

import os
import docx
import numpy as np
import pandas as pd
from pandas import Series,DataFrame
import sys

source_docxs = []
# In[2]指定源docx文档目录
file_path = r"D:\_PyCharmWorkSpace\_homework_interesting\1717测绘毕业要求评价表_source_2"
for file in os.listdir(file_path):
    if file.endswith(".docx") or file.endswith(".doc"):
        source_docxs.append(file_path + "\\" + file)

file = source_docxs[0]

docx_file = docx.Document(file)

row_str = []
for table in docx_file.tables:
    for row in table.rows:
        row_str.append([cell.text for cell in row.cells])
row_str = np.array(row_str)

# 这里应当是需要修改的，2:14是在我知道行列的情况下
# 实际应使用的是row_str[0], row_str[1]来代替14和6    #这个已更正好了
# 2和1不需要代替，是因为需要指定哪几行是行标题和列标题
frame = DataFrame(row_str)
# frame.iloc[2:14,1:6]=0
frame.iloc[2:row_str.shape[0],1:row_str.shape[1]]=0

for i in range(len(source_docxs)):
    file = source_docxs[i]
    
    docx_file = docx.Document(file)
    
    docx_table = docx_file.tables[0]
    
    for table in docx_file.tables:
        for i in range(2, row_str.shape[0]):
            for j in range(1, row_str.shape[1]):
                table.cell(i,j).text = table.cell(i,j).text + "\\"
    
    for table in docx_file.tables:
        for i in range(2, row_str.shape[0]):
            for j in range(1, row_str.shape[1]):
                if str(table.cell(i,j).text) == "\\":
                    table.cell(i,j).text = "0"
                else:
                    table.cell(i,j).text = "1"
                    
    row_str = []
    for table in docx_file.tables:
        for row in table.rows:
            row_str.append([cell.text for cell in row.cells])

    row_str = np.array(row_str)
#     print("rows："+str(row_str.shape[0]))
#     print("columns："+str(row_str.shape[1]))
    
    row_str_pd = pd.DataFrame()
    row_str_pd = pd.DataFrame(row_str)
    
    for i in range(1,row_str.shape[1]):
        row_str_pd.iloc[2:row_str.shape[0], i] = pd.to_numeric(row_str_pd.iloc[2:row_str.shape[0], i])
    
    frame.iloc[2:row_str.shape[0],1:row_str.shape[1]] += row_str_pd.iloc[2:row_str.shape[0],1:row_str.shape[1]]
    
    sys.stdout.write("\r"+ "文件编号"+ str(i) + file)
    sys.stdout.flush()

new_docx = docx.Document()

new_docx_table = new_docx.add_table(rows = row_str.shape[0], cols = row_str.shape[1])

for i in range(2, row_str.shape[0]):
        for j in range(1, row_str.shape[1]):
            new_docx_table.cell(i,j).text = "0"

for table in new_docx.tables:
    # 有几行列标题，应手动输入
    for i in range(2):
        for j in range(row_str.shape[1]):
            table.cell(i,j).text = row_str[i,j]
for table in new_docx.tables:
    # 从第几行开始设置行标题，应手动输入
    for i in range(2,row_str.shape[0]):
        for j in range(1):
            table.cell(i,j).text = row_str[i,j]

for i in range(2, row_str.shape[0]):
    for j in range(1, row_str.shape[1]):
        new_docx_table.cell(i,j).text = str(frame.iloc[i,j])

new_docx.save("new_docx.docx")
